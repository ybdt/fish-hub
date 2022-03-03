#!/usr/bin/env python3

'''
由于公司安排的批量操作word的工作太过繁琐，故写下此半自动化工具
参考链接：
https://blog.csdn.net/woshisangsang/article/details/75221723
https://blog.csdn.net/zhengyikuangge/article/details/80451424
https://zhuanlan.zhihu.com/p/90855359
https://www.jianshu.com/p/7d2fcf976914
'''

import docx
import os

def duplicate_word_file():
    f_r = open("2019年6月1周.docx", "rb");
    content = f_r.read();

    for year in range(2019, 2021):
        if year == 2019:
            for month in range(6, 13):
                for week in range(1, 5):
                    name = str(year) + "年" + str(month) + "月" + str(week) + "周.docx";
                    f_w = open(name, "wb");
                    f_w.write(content);
        else:
            for month in range(1, 13):
                for week in range(1, 5):
                    name = str(year) + "年" + str(month) + "月" + str(week) + "周.docx";
                    f_w = open(name, "wb");
                    f_w.write(content);
    
    f_r.close();
    f_w.close();

def extract_word_content():
    file = docx.Document("运维事件单2019.6.4.docx");
    print( "段落数：" + str(len(file.paragraphs)) );
    print(file.paragraphs[0].text);

def extract_word_table_content():
    f = open("extract_word_table_content.txt", "w");
    current_items = os.listdir();
    
    for item in current_items:
        if item.endswith(".docx"):
            file = docx.Document(item);
            all_tables = file.tables;
            #print(len(all_tables));
            first_table = all_tables[0];
            rows = len(first_table.rows);
            for i in range(0, rows):
                if i == 5:
                    word_content = first_table.cell(i, 1).text;
                    print(word_content);
                    item_content = item.strip("运维事件单").strip(".docx");
                    print(item_content);
                    f.write(item_content + ": " + word_content + "\n");
    
    f.close();

def replace_word_content():
    file = docx.Document("2019年6月1周.docx");
    print( len(file.paragraphs) );#55
    print(file.paragraphs[18].text);#'检测日期：2019年6月1周'
    '''
    for i in range(0, 55):
        if "2019年6月1周" in file.paragraphs[i].text:
             print(i);#18
    '''
    print(file.paragraphs[18].text);
    print( len(file.paragraphs[18].runs) );#7
    for i in file.paragraphs[18].runs:
        print(i.text);
    file.save("a.docx");

def main():
    #duplicate_word_file();
    #extract_word_content();
    extract_word_table_content();

if __name__ == "__main__":
    main();
